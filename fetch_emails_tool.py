from smolagents.tools import tool
import os
import base64
import pickle
from pathlib import Path
from email import message_from_bytes
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import fitz  # PyMuPDF
from bs4 import BeautifulSoup  # <-- Make sure to install with: pip install beautifulsoup4
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def convert_docx_to_pdf_text(docx_path, pdf_path):
    doc = Document(docx_path)
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    y = 750
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            c.drawString(30, y, text)
            y -= 15
            if y < 50:
                c.showPage()
                y = 750
    c.save()


def save_email_text_to_pdf(text, pdf_path):
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter,
                            rightMargin=30, leftMargin=30,
                            topMargin=30, bottomMargin=30)

    styles = getSampleStyleSheet()
    story = []

    for line in text.splitlines():
        line = line.strip()
        if line:
            para = Paragraph(line, styles["Normal"])
            story.append(para)
            story.append(Spacer(1, 12))

    doc.build(story)


def convert_docx_to_wrapped_pdf(docx_path, output_pdf_path):
    doc = Document(docx_path)
    doc_template = SimpleDocTemplate(str(output_pdf_path), pagesize=letter,
                                     rightMargin=30, leftMargin=30,
                                     topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            p = Paragraph(text, styles["Normal"])
            story.append(p)
            story.append(Spacer(1, 12))

    doc_template.build(story)


    
@tool
def fetch_recent_emails(k: int = 3) -> str:
    """
    Fetches the last k emails from Gmail, extracts the message body (preferring plain text),
    and saves each email as a merged PDF in the 'emails/' folder along with its PDF attachments.
    Do not try to save the files after calling this tool. No saving tool must be used after this tool.

    Args:
        k (int): Number of recent emails to fetch.

    Returns:
        str: Message indicating the result of the operation.
    """
    SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
    creds = None

    # Load credentials
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)

    service = build('gmail', 'v1', credentials=creds)
    results = service.users().messages().list(userId='me', maxResults=k).execute()
    messages = results.get('messages', [])

    email_dir = Path("emails")
    email_dir.mkdir(exist_ok=True)

    for idx, msg in enumerate(messages):
        msg_data = service.users().messages().get(userId='me', id=msg['id'], format='raw').execute()
        raw_msg = base64.urlsafe_b64decode(msg_data['raw'].encode('UTF-8'))
        mime_msg = message_from_bytes(raw_msg)

        # Prepare text content
        subject = mime_msg['subject'] or "(No Subject)"
        sender = mime_msg['from'] or "(Unknown Sender)"
        email_text = f"Subject: {subject}\nFrom: {sender}\n\n"
        attachments = []
        email_body_found = False

        if mime_msg.is_multipart():
            for part in mime_msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition", ""))

                if content_type == "text/plain" and "attachment" not in content_disposition:
                    email_text += part.get_payload(decode=True).decode(errors='ignore')
                    email_body_found = True

                elif content_type == "text/html" and "attachment" not in content_disposition and not email_body_found:
                    html = part.get_payload(decode=True).decode(errors='ignore')
                    soup = BeautifulSoup(html, 'html.parser')
                    email_text += soup.get_text(separator="\n", strip=True)
                    email_body_found = True

                elif "attachment" in content_disposition:
                    filename = part.get_filename()
                    if filename and filename.lower().endswith((".pdf", ".doc", ".docx")):
                        file_path = email_dir / f"{idx+1}_{filename}"
                        with open(file_path, "wb") as f:
                            f.write(part.get_payload(decode=True))
                        attachments.append(file_path)
        else:
            payload = mime_msg.get_payload(decode=True).decode(errors='ignore')
            try:
                soup = BeautifulSoup(payload, 'html.parser')
                email_text += soup.get_text(separator="\n", strip=True)
            except Exception:
                email_text += payload

        # Save plain text to a temporary PDF
        text_pdf = email_dir / f"{idx+1}_text.pdf"
        save_email_text_to_pdf(email_text, text_pdf)

        # Merge into final PDF
        merged_pdf_path = email_dir / f"email_{idx+1}.pdf"
        pdf_writer = fitz.open()
        pdf_writer.insert_pdf(fitz.open(str(text_pdf)))

        for attachment in attachments:
            try:
                ext = attachment.suffix.lower()
                if ext == ".pdf":
                    pdf_writer.insert_pdf(fitz.open(str(attachment)))

                elif ext == ".docx":
                    docx_pdf = attachment.with_suffix(".converted.pdf")
                    convert_docx_to_wrapped_pdf(attachment, docx_pdf)
                    pdf_writer.insert_pdf(fitz.open(str(docx_pdf)))
                    docx_pdf.unlink()

            except Exception as e:
                print(f"Skipping {attachment.name}: {e}")
                continue

        pdf_writer.save(str(merged_pdf_path))
        pdf_writer.close()

        # Clean up intermediate files
        text_pdf.unlink()
        for attachment in attachments:
            attachment.unlink()

    return f"{k} emails saved as merged PDFs in {email_dir.resolve()}"
